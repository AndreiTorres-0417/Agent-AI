import json
import logging
import os
import re
from io import BytesIO
from typing import Any, Dict, Generator, List, Optional

import requests
from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import JSONResponse
from fastapi.responses import FileResponse
from fastapi.responses import StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from dotenv import load_dotenv


load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("academic-review-mvp")

app = FastAPI(title="Academic Review MVP", version="0.1.0")

if os.path.isdir("static"):
    app.mount("/static", StaticFiles(directory="static"), name="static")


@app.exception_handler(Exception)
async def unhandled_exception_handler(request: Request, exc: Exception) -> JSONResponse:
    logger.exception("Unhandled API error at %s", request.url.path)
    return JSONResponse(
        status_code=500,
        content={
            "detail": "Unhandled server error",
            "error_type": exc.__class__.__name__,
            "error": str(exc),
        },
    )


class ChatRequest(BaseModel):
    session_id: str = Field(..., description="Client-provided chat session ID")
    message: str = Field(..., min_length=1, description="Latest user message")
    model: Optional[str] = Field(None, description="Academic assistant model: OpenAI or local Gemma")


class ChatResponse(BaseModel):
    session_id: str
    phase: str
    questions: List[str]
    next_prompt: str
    context: Dict[str, Any]
    analysis: Optional[Dict[str, Any]] = None


class AnalyzeTextRequest(BaseModel):
    text: str = Field(..., min_length=20)
    session_id: Optional[str] = Field(None, description="Chat session that should receive the paper context")
    metadata: Optional[Dict[str, Any]] = None
    model: Optional[str] = Field(None, description="Document analysis model: OpenAI or local Gemma")
    review_mode: Optional[str] = Field("academic", description="Review mode: academic")
    format_mode: Optional[str] = Field("ieee", description="Compliance style: ieee")


class AnalyzeTextStreamRequest(AnalyzeTextRequest):
    pass


class IssueItem(BaseModel):
    issue: str
    severity: str
    evidence: str
    recommendation: str


class SuggestionItem(BaseModel):
    priority: str
    suggestion: str
    rationale: str
    expected_impact: str


class RewriteItem(BaseModel):
    original_excerpt: str
    rewritten_excerpt: str
    reason: str


class HighlightItem(BaseModel):
    excerpt: str
    message: str
    severity: str
    category: str


class FormatChangeItem(BaseModel):
    issue: str
    severity: str
    category: str = "IEEE formatting"
    original: str
    replacement: str
    note: str
    sub_issues: List[str] = Field(default_factory=list)


class FormatTextRequest(BaseModel):
    text: str = Field(..., min_length=20)


class ExportDocxRequest(BaseModel):
    text: str = Field(..., min_length=1)
    filename: str = "ieee_formatted.docx"


class FormatResponse(BaseModel):
    summary: str
    fixed_text: str
    changes: List[FormatChangeItem]
    highlights: List[HighlightItem]
    transformations_applied: bool = False
    transformation_count: int = 0
    source: str = "deterministic_ieee_formatter"


class AnalysisResponse(BaseModel):
    summary: str
    structure_format_issues: List[IssueItem]
    academic_quality_issues: List[IssueItem]
    citation_consistency_issues: List[IssueItem]
    prioritized_suggestions: List[SuggestionItem]
    optional_rewrite_suggestions: List[RewriteItem]
    highlights: List[HighlightItem]
    reviewed_text: Optional[str] = None
    source: str
    fallback_used: bool


chat_sessions: Dict[str, Dict[str, Any]] = {}


@app.get("/", include_in_schema=False)
def index() -> FileResponse:
    return FileResponse("static/index.html")


ALLOWED_MODELS = {"openai:gpt-4.1-mini", "gemma3:1b"}
ALLOWED_REVIEW_MODES = {"format", "academic"}
ALLOWED_FORMAT_MODES = {"apa7", "ieee"}
REFERENCE_HEADING_RE = re.compile(r"^\s*#{0,6}\s*references\s*$", re.IGNORECASE)


def get_analysis_model(model_override: Optional[str] = None) -> str:
    model = model_override or os.getenv("ANALYSIS_MODEL", "openai:gpt-4.1-mini")
    if model not in ALLOWED_MODELS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported model '{model}'. Choose one of: {', '.join(sorted(ALLOWED_MODELS))}.",
        )
    return model


def get_ollama_url() -> str:
    return os.getenv("OLLAMA_URL", "http://localhost:11434").rstrip("/")


def get_ollama_timeout() -> int:
    try:
        return int(os.getenv("OLLAMA_TIMEOUT_SECONDS", "600"))
    except ValueError:
        return 600


def ollama_options(kind: str = "analysis") -> Dict[str, Any]:
    num_predict_by_kind = {
        "chat": 300,
        "analysis": 2200,
    }
    return {
        "temperature": 0.2 if kind != "chat" else 0.3,
        "num_predict": num_predict_by_kind.get(kind, 2200),
    }


def is_openai_model(model: str) -> bool:
    return model.startswith("openai:")


def openai_model_id(model: str) -> str:
    return model.removeprefix("openai:")


def get_openai_api_key() -> str:
    key = os.getenv("OPENAI_API_KEY", "").strip()
    if not key:
        raise RuntimeError("OPENAI_API_KEY is not configured. Add it to .env to use OpenAI hosted models.")
    return key


ACADEMIC_ASSISTANT_PROMPT = """
You are an academic-paper assistant. Help students plan, improve, and understand academic papers.
Provide practical guidance for outlines, thesis statements, section organization, literature synthesis,
citations, references, tables, figures, revisions, and APA 7 or IEEE compliance.
Do not invent sources, quotations, data, DOI values, or page numbers. State clearly when a source or fact
must be verified. Do not write an entire submit-ready paper for the student; help them reason, revise, and
produce their own work.

Keep responses concise and direct. For greetings or broad requests for help, reply with a short description
of what you can help with and ask the user to share their paper, question, or goal. Do not start with a
questionnaire or a numbered list of intake questions. Ask at most one focused follow-up question when the
user's specific request cannot be answered without clarification. Do not use emojis or excessive encouragement.
""".strip()


SECTION_ALIASES = {
    "abstract": ["abstract"],
    "introduction": ["introduction", "intro"],
    "literature review": ["literature review", "related work", "background"],
    "methodology": ["methodology", "methods", "method", "materials and methods"],
    "results": ["results", "findings"],
    "discussion": ["discussion"],
    "conclusion": ["conclusion", "conclusions"],
    "references": ["references", "reference list", "bibliography", "works cited"],
}


def normalize_section_title(title: str) -> str:
    cleaned = re.sub(r"^\s*#{1,6}\s*", "", title).strip()
    cleaned = re.sub(r"^(?:[IVXLC]+|\d+)\.\s*", "", cleaned, flags=re.IGNORECASE).strip()
    return re.sub(r"\s+", " ", cleaned).lower()


def looks_like_section_heading(line: str) -> Optional[str]:
    stripped = line.strip()
    if not stripped:
        return None
    markdown_heading = re.match(r"^#{1,6}\s+(.+?)\s*$", stripped)
    if markdown_heading:
        return markdown_heading.group(1).strip()
    numbered_heading = re.match(r"^(?:[IVXLC]+|\d+)\.\s+(.+?)\s*$", stripped, flags=re.IGNORECASE)
    if numbered_heading:
        return numbered_heading.group(1).strip()
    normalized = normalize_section_title(stripped)
    known_names = {alias for aliases in SECTION_ALIASES.values() for alias in aliases}
    if normalized in known_names:
        return stripped
    return None


def extract_paper_sections(text: str) -> List[Dict[str, str]]:
    lines = text.splitlines()
    headings: List[Dict[str, Any]] = []
    for idx, line in enumerate(lines):
        title = looks_like_section_heading(line)
        if title:
            headings.append({"line": idx, "title": title, "heading": line.strip()})
    sections: List[Dict[str, str]] = []
    for pos, heading in enumerate(headings):
        start = int(heading["line"]) + 1
        end = int(headings[pos + 1]["line"]) if pos + 1 < len(headings) else len(lines)
        content = "\n".join(lines[start:end]).strip()
        if content:
            sections.append(
                {
                    "title": str(heading["title"]),
                    "normalized_title": normalize_section_title(str(heading["title"])),
                    "heading": str(heading["heading"]),
                    "content": content,
                }
            )
    return sections


def requested_section_key(message: str) -> Optional[str]:
    lowered = message.lower()
    for key, aliases in SECTION_ALIASES.items():
        for alias in aliases:
            if re.search(rf"\b{re.escape(alias)}\b", lowered):
                return key
    return None


def find_requested_section(message: str, sections: List[Dict[str, str]]) -> Optional[Dict[str, str]]:
    key = requested_section_key(message)
    if not key:
        return None
    aliases = SECTION_ALIASES[key]
    for section in sections:
        title = section["normalized_title"]
        if title == key or title in aliases or any(alias in title for alias in aliases):
            return section
    return {
        "title": key.title(),
        "normalized_title": key,
        "heading": key.title(),
        "content": "",
    }


def call_assistant_chat(
    messages: List[Dict[str, str]],
    model_override: Optional[str] = None,
    paper_context: Optional[str] = None,
    focused_section: Optional[Dict[str, str]] = None,
) -> str:
    model = model_override or os.getenv("OPENAI_CHAT_MODEL", "openai:gpt-4.1-mini")
    if model not in ALLOWED_MODELS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported model '{model}'. Choose one of: {', '.join(sorted(ALLOWED_MODELS))}.",
        )
    system_prompt = ACADEMIC_ASSISTANT_PROMPT
    if paper_context:
        system_prompt += (
            "\n\nThe user has analyzed a paper in Step 2. Use the following paper context when answering "
            "follow-up questions. Do not claim details beyond this context:\n" + paper_context
        )
    if focused_section:
        if focused_section.get("content"):
            system_prompt += (
                "\n\nThe latest user question appears to ask about this specific paper section. "
                "Analyze this section directly and do not guess about section content outside this excerpt unless needed for brief context.\n"
                f"Requested section: {focused_section.get('title', '')}\n"
                f"Section text:\n{focused_section.get('content', '')[:5000]}"
            )
        else:
            system_prompt += (
                "\n\nThe latest user question appears to ask about a section, but that section was not found in the analyzed paper. "
                f"Requested section: {focused_section.get('title', '')}. Say that the section was not found and ask the user to paste it."
            )
    chat_messages = [{"role": "system", "content": system_prompt}, *messages]
    if not is_openai_model(model):
        resp = requests.post(
            f"{get_ollama_url()}/api/chat",
            json={
                "model": model,
                "messages": chat_messages,
                "stream": False,
                "options": ollama_options("chat"),
            },
            timeout=get_ollama_timeout(),
        )
        if resp.status_code >= 400:
            raise RuntimeError(f"Ollama request failed with HTTP {resp.status_code}: {resp.text.strip()}")
        return resp.json()["message"]["content"].strip()

    resp = requests.post(
        "https://api.openai.com/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {get_openai_api_key()}",
            "Content-Type": "application/json",
        },
        json={
            "model": openai_model_id(model),
            "messages": chat_messages,
            "temperature": 0.3,
            "max_tokens": 900,
        },
        timeout=90,
    )
    if resp.status_code >= 400:
        raise RuntimeError(f"OpenAI request failed with HTTP {resp.status_code}: {resp.text.strip()}")
    return resp.json()["choices"][0]["message"]["content"].strip()


def normalize_format_mode(format_mode: Optional[str] = None) -> str:
    mode = (format_mode or "ieee").lower()
    if mode not in ALLOWED_FORMAT_MODES:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported format mode '{format_mode}'. Choose one of: {', '.join(sorted(ALLOWED_FORMAT_MODES))}.",
        )
    return mode


def normalize_review_mode(review_mode: Optional[str] = None) -> str:
    mode = (review_mode or "academic").lower()
    if mode not in ALLOWED_REVIEW_MODES:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported review mode '{review_mode}'. Choose one of: {', '.join(sorted(ALLOWED_REVIEW_MODES))}.",
        )
    return mode


def int_to_roman(value: int) -> str:
    pairs = [
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    ]
    result = ""
    for number, roman in pairs:
        while value >= number:
            result += roman
            value -= number
    return result


def reference_number_from_text(value: str) -> Optional[str]:
    match = re.search(r"\[(\d+)\]", value)
    return match.group(1) if match else None


def clean_ieee_findings(changes: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    grouped: Dict[str, Dict[str, Any]] = {}
    allowed_categories = {
        "IEEE formatting",
        "IEEE citation issues",
        "citation consistency",
        "citation numbering issues",
        "research methodology",
        "academic tone",
        "reference heading",
        "figure label",
        "table label",
    }

    def severity_rank(value: str) -> int:
        return {"low": 1, "medium": 2, "high": 3}.get(value, 2)

    def add_group(
        key: str,
        issue: str,
        severity: str,
        category: str,
        original: str,
        replacement: str,
        note: str,
        sub_issue: Optional[str] = None,
    ) -> None:
        current = grouped.get(key)
        if not current:
            grouped[key] = {
                "issue": issue,
                "severity": severity,
                "category": category,
                "original": original,
                "replacement": replacement,
                "note": note,
                "sub_issues": [],
            }
            current = grouped[key]
        elif severity_rank(severity) > severity_rank(current["severity"]):
            current["severity"] = severity
        if sub_issue and sub_issue not in current["sub_issues"]:
            current["sub_issues"].append(sub_issue)

    for item in changes:
        issue = item.get("issue", "")
        category = item.get("category", "IEEE formatting")
        if category not in allowed_categories:
            continue
        original = item.get("original", "")
        replacement = item.get("replacement", "")
        note = item.get("note", "")
        severity = item.get("severity", "medium")
        ref_no = reference_number_from_text(original) or reference_number_from_text(replacement)

        if ref_no and (
            "APA-style" in issue
            or "author order" in issue
            or "ampersand" in issue
            or "year placement" in issue
            or "multiple-author" in issue
        ):
            add_group(
                f"ref-{ref_no}-ieee-reference",
                f"Reference [{ref_no}] does not comply with IEEE reference format",
                "high",
                "IEEE citation issues",
                original,
                "Rewrite using IEEE reference structure, e.g., [n] J. Smith, \"Article Title,\" Source, year.",
                "Multiple formatting patterns were detected in the same reference entry.",
                issue,
            )
            continue

        if ref_no and (
            "incomplete" in issue.lower()
            or "raw url" in issue.lower()
            or "web reference" in issue.lower()
            or "et al." in issue.lower()
            or "quotation marks" in issue.lower()
        ):
            add_group(
                f"ref-{ref_no}-ieee-reference",
                f"Reference [{ref_no}] does not comply with IEEE reference format",
                severity,
                "IEEE citation issues",
                original,
                "Add complete IEEE metadata: initials and surnames, quoted title, source/site, year, URL/DOI if applicable, and access date for web sources.",
                "The reference is missing required IEEE metadata or structure.",
                issue,
            )
            continue

        if "Citation order" in issue or "skipped" in issue.lower() or "not tied" in issue:
            add_group(
                "citation-numbering-sync",
                "Citation numbering and reference list are not synchronized",
                "high",
                "citation numbering issues",
                original,
                "Review citation order manually; do not renumber automatically unless each source identity is verified.",
                "IEEE numbering must follow first appearance and every listed reference should match an in-text citation.",
                issue,
            )
            continue

        if "Mixed citation systems" in issue or "Narrative source mention" in issue:
            add_group(
                "mixed-citation-systems",
                "The paper mixes IEEE numbered citations with author-name citation style",
                "high",
                "citation consistency",
                original,
                "Use numbered IEEE citations consistently, e.g., replace author-name source mentions with the correct bracket number after verifying the source.",
                "The text contains both numbered citations and author-name references.",
                issue,
            )
            continue

        if "causal" in issue.lower() or "overstated" in issue.lower():
            add_group(
                "survey-causation",
                "Survey-based methodology does not support causal conclusions",
                "high",
                "research methodology",
                original,
                "Use wording such as 'is associated with' or 'shows a relationship with.'",
                "The paper uses causal language even though the method appears to be a survey.",
                issue,
            )
            continue

        key = f"{category}-{issue}-{original[:40]}"
        add_group(key, issue, severity, category, original, replacement, note)

    cleaned = list(grouped.values())
    cleaned.sort(key=lambda item: ({"high": 0, "medium": 1, "low": 2}.get(item["severity"], 1), item["category"], item["issue"]))
    return cleaned


def split_applied_transformations(changes: List[Dict[str, Any]], fixed_text: str) -> tuple[List[Dict[str, Any]], int]:
    transformation_keywords = [
        "converted",
        "changed",
        "normalized",
        "numbered",
        "revised",
    ]
    unresolved: List[Dict[str, Any]] = []
    transformation_count = 0
    allowed_categories = {
        "IEEE formatting",
        "IEEE citation issues",
        "citation consistency",
        "citation numbering issues",
        "research methodology",
        "academic tone",
        "reference heading",
        "figure label",
        "table label",
    }

    for item in changes:
        if item.get("category", "IEEE formatting") not in allowed_categories:
            continue
        if item.get("category") in {"academic tone", "research methodology"}:
            unresolved.append(item)
            continue
        issue = item.get("issue", "").lower()
        original = item.get("original", "")
        looks_like_applied_fix = any(keyword in issue for keyword in transformation_keywords)

        if looks_like_applied_fix and (not original or original not in fixed_text):
            transformation_count += 1
            continue

        if item.get("issue") == "Reference entry numbered" and re.match(r"^\[\d+\]", item.get("replacement", "")):
            transformation_count += 1
            continue

        unresolved.append(item)

    return unresolved, transformation_count


def ieee_format_text(text: str) -> Dict[str, Any]:
    fixed = text
    changes: List[Dict[str, str]] = []
    highlights: List[Dict[str, str]] = []
    citation_map: Dict[str, int] = {}

    def add_change(issue: str, severity: str, original: str, replacement: str, note: str, category: str = "IEEE formatting") -> None:
        changes.append(
            {
                "issue": issue,
                "severity": severity,
                "category": category,
                "original": original,
                "replacement": replacement,
                "note": note,
            }
        )
        if original:
            highlights.append(
                {
                    "excerpt": original,
                    "message": note,
                    "severity": severity,
                    "category": category,
                }
            )

    def citation_number(key: str) -> int:
        if key not in citation_map:
            citation_map[key] = len(citation_map) + 1
        return citation_map[key]

    def replace_apa_citation(match: re.Match[str]) -> str:
        original = match.group(0)
        author = match.group(1).strip()
        year = match.group(2).strip()
        number = citation_number(f"{author.lower()}-{year}")
        replacement = f"[{number}]"
        add_change(
            "APA author-date citation converted to IEEE numeric citation",
            "high",
            original,
            replacement,
            "IEEE uses numbered bracket citations in order of first appearance.",
                    "IEEE citation issues",
                )
        return replacement

    fixed = re.sub(r"\(([A-Z][A-Za-z\s&.-]{1,60}),\s*(\d{4}[a-z]?)\)", replace_apa_citation, fixed)

    def replace_bare_author_citation(match: re.Match[str]) -> str:
        original = match.group(0)
        author = match.group(1).strip()
        number = citation_number(author.lower())
        replacement = f"[{number}]"
        add_change(
            "Bare author parenthetical citation converted to IEEE numeric citation",
            "high",
            original,
            replacement,
            "IEEE does not use parenthetical author-only citations such as (Smith). Use a numbered citation.",
            "in-text citation",
        )
        return replacement

    fixed = re.sub(r"\(([A-Z][A-Za-z.-]{2,30}(?:\s+et\s+al\.)?)\)", replace_bare_author_citation, fixed)

    def replace_author_year_bracket(match: re.Match[str]) -> str:
        original = match.group(0)
        author = match.group(1).strip()
        year = match.group(2).strip()
        number = citation_number(f"{author.lower()}-{year}")
        replacement = f"[{number}]"
        add_change(
            "Non-IEEE author-year bracket citation converted",
            "high",
            original,
            replacement,
            "IEEE does not use author [year] citation style.",
            "IEEE citation issues",
        )
        return replacement

    fixed = re.sub(r"\b([A-Z][A-Za-z.-]{2,30})\s*\[(\d{4}[a-z]?)\]", replace_author_year_bracket, fixed)

    narrative_uncited = re.findall(
        r"\b([A-Z][A-Za-z.-]{2,30}(?:\s+et\s+al\.)?)\s+(?:argues?|states?|found|finds|suggests?|indicates?|demonstrates?|claims?|reports?|concludes?)\b(?!\s*\[\d+\])",
        fixed,
    )
    trailing_author_after_citation = re.findall(
        r"(?:\[\d+\](?:,\s*)?)+(?:\s*,?\s*and\s+|\s*,\s*)([A-Z][A-Za-z.-]{2,30}(?:\s+et\s+al\.?)?)",
        fixed,
    )
    narrative_uncited.extend(trailing_author_after_citation)
    for author in sorted(set(narrative_uncited), key=str.lower):
        add_change(
            "Narrative source mention may be missing IEEE bracket citation",
            "high",
            author,
            f"{author} [n]",
            "IEEE narrative citations still need a numbered bracket citation, e.g., Smith [1] states...",
            "in-text citation",
        )

    has_author_parenthetical = bool(re.search(r"\([A-Z][A-Za-z.-]{2,30}(?:,\s*\d{4})?\)", text))
    has_author_year_bracket = bool(re.search(r"\b[A-Z][A-Za-z.-]{2,30}\s*\[\d{4}\]", text))
    has_ieee_brackets = bool(re.search(r"\[\d+\]", text))
    has_narrative_source = bool(narrative_uncited)
    mixed_count = sum([has_author_parenthetical, has_author_year_bracket, has_ieee_brackets, has_narrative_source])
    if mixed_count >= 2:
        add_change(
            "Mixed citation systems detected",
            "high",
            "mixed citation styles",
            "IEEE numbered citations only",
            "The draft appears to mix author-date/author-only/narrative citations with IEEE bracket citations.",
            "citation consistency",
        )

    survey_design = bool(re.search(r"\bsurvey\b|\bquestionnaire\b|\brespondents?\b", fixed, re.IGNORECASE))
    causal_overclaim = re.findall(
        r"\b(directly causes?|causes?|determine whether .{0,80}? causes?|directly improves?|causes better .{0,40})\b",
        fixed,
        re.IGNORECASE,
    )
    if survey_design and causal_overclaim:
        add_change(
            "Causal claim is not supported by survey design",
            "high",
            causal_overclaim[0],
            "is associated with",
            "A survey can usually support association, not direct causation.",
            "research methodology",
        )

    overstated_findings = re.findall(
        r"\b(coffee directly improves [^.]+|coffee consumption causes [^.]+|directly improves [^.]+|causes better [^.]+)\b",
        fixed,
        re.IGNORECASE,
    )
    for phrase in sorted(set(overstated_findings), key=str.lower):
        add_change(
            "Overstated causal finding",
            "high",
            phrase,
            phrase.replace("directly improves", "is associated with").replace("causes", "is associated with"),
            "Survey results should be framed as association or relationship unless the design supports causality.",
            "research methodology",
        )

    for wrong_heading in ["Bibliography", "Works Cited", "Reference List"]:
        pattern = re.compile(rf"(?im)^\s*{re.escape(wrong_heading)}\s*$")
        if pattern.search(fixed):
            fixed = pattern.sub("References", fixed)
            add_change(
                "Reference list title changed to IEEE style",
                "medium",
                wrong_heading,
                "References",
                "IEEE papers use the heading 'References'.",
                "reference heading",
            )

    fixed = re.sub(
        r"(?im)^Figure\s+(\d+)\s*[:.-]\s*(.+)$",
        lambda m: f"Fig. {m.group(1)}. {m.group(2).strip()}",
        fixed,
    )
    if re.search(r"(?im)^Figure\s+\d+\s*[:.-]", text):
        add_change(
            "Figure label normalized",
            "medium",
            "Figure n:",
            "Fig. n.",
            "IEEE figure captions use 'Fig. 1. Caption text.'",
            "figure label",
        )

    fixed = re.sub(
        r"(?im)^Table\s+(\d+)\s*[:.-]\s*(.+)$",
        lambda m: f"Table {int_to_roman(int(m.group(1)))}. {m.group(2).strip()}",
        fixed,
    )
    if re.search(r"(?im)^Table\s+\d+\s*[:.-]", text):
        add_change(
            "Table label normalized",
            "medium",
            "Table 1:",
            "Table I.",
            "IEEE table captions use Roman numerals such as 'Table I.'",
            "table label",
        )

    tone_replacements: List[tuple[str, str, str]] = [
        (r"\bproves\b", "suggests", "IEEE/academic reporting should avoid 'proves' when the method only supports association."),
        (r"\bdefinitely\b", "strongly suggests", "Use cautious academic wording."),
        (r"\bvery big impact\b", "substantial association", "Use precise academic phrasing."),
        (r"\ba lot of studies say\b", "prior studies suggest", "Use formal academic wording."),
    ]
    for pattern, replacement, note in tone_replacements:
        if re.search(pattern, fixed, flags=re.IGNORECASE):
            original = re.search(pattern, fixed, flags=re.IGNORECASE).group(0)
            fixed = re.sub(pattern, replacement, fixed, flags=re.IGNORECASE)
            add_change("Informal or overstrong wording revised", "medium", original, replacement, note, "academic tone")

    lines = fixed.splitlines()
    ref_start = None
    for idx, line in enumerate(lines):
        if re.match(r"^\s*References\s*$", line, re.IGNORECASE):
            ref_start = idx
            break

    if ref_start is not None:
        body_text = "\n".join(lines[:ref_start])
        body_citations = [int(x) for x in re.findall(r"\[(\d+)\]", body_text)]
        body_citation_set = set(body_citations)
        reference_numbers_seen: List[int] = []
        ref_number = 1
        for idx in range(ref_start + 1, len(lines)):
            line = lines[idx].strip()
            if not line:
                continue
            ref_match_before = re.match(r"^\[(\d+)\]", line)
            if ref_match_before:
                reference_numbers_seen.append(int(ref_match_before.group(1)))
            if not re.match(r"^\[\d+\]", line):
                lines[idx] = f"[{ref_number}] {line}"
                add_change(
                    "Reference entry numbered",
                    "high",
                    line[:80],
                    lines[idx][:80],
                    "Each IEEE reference entry should start with a bracketed number.",
                    "IEEE citation issues",
                )
            else:
                current = re.match(r"^\[(\d+)\]", line)
                if current and int(current.group(1)) != ref_number:
                    new_line = re.sub(r"^\[\d+\]", f"[{ref_number}]", line)
                    add_change(
                        "Reference number order corrected",
                        "high",
                        line[:80],
                        new_line[:80],
                        "IEEE references should be numbered in citation order.",
                        "citation numbering issues",
                    )
                    lines[idx] = new_line
            author_order_match = re.match(r"^(\[\d+\]\s+)([A-Z][A-Za-z-]+),\s+([A-Z])\.\s*(.*)$", lines[idx])
            if author_order_match:
                original_line = lines[idx]
                lines[idx] = (
                    f"{author_order_match.group(1)}"
                    f"{author_order_match.group(3)}. {author_order_match.group(2)}, "
                    f"{author_order_match.group(4)}"
                )
                add_change(
                    "Reference author order converted to IEEE initials-before-surname format",
                    "high",
                    original_line[:80],
                    lines[idx][:80],
                    "A safe author-order transformation was applied for a single surname/initial pattern.",
                    "IEEE citation issues",
                )
            if "&" in lines[idx]:
                original_line = lines[idx]
                lines[idx] = lines[idx].replace(" & ", " and ")
                add_change(
                    "APA ampersand converted in reference entry",
                    "medium",
                    original_line[:80],
                    lines[idx][:80],
                    "IEEE references should not use APA-style ampersands between authors.",
                    "IEEE citation issues",
                )
            if re.match(r"^[A-Z][a-z]+,\s+[A-Z][a-z]+", re.sub(r"^\[\d+\]\s*", "", lines[idx])):
                add_change(
                    "Author name may not follow IEEE initials-before-surname format",
                    "medium",
                    lines[idx][:80],
                    lines[idx][:80],
                    "IEEE references usually use initials before surname, e.g., J. Smith.",
                    "IEEE citation issues",
                )
            if re.search(r"^\[\d+\]\s+[A-Z][A-Za-z-]+,\s+[A-Z]\.", lines[idx]):
                add_change(
                    "Reference author order is APA-like, not IEEE",
                    "high",
                    lines[idx][:80],
                    "Use initials before surname, e.g., [1] J. Smith, ...",
                    "IEEE uses initials before surname, e.g., J. Smith, not Smith, J.",
                    "IEEE citation issues",
                )
            if re.search(r"\bet\s+al\.?,", lines[idx], re.IGNORECASE):
                add_change(
                    "Reference entry uses incomplete 'et al.' author listing",
                    "high",
                    lines[idx][:80],
                    lines[idx][:80],
                    "IEEE references should provide the available author names in the required reference format, not a vague 'et al.' placeholder.",
                    "IEEE citation issues",
                )
            if re.search(r"\(\d{4}\)", lines[idx]):
                add_change(
                    "APA-style year placement detected in IEEE reference",
                    "high",
                    lines[idx][:80],
                    lines[idx][:80],
                    "IEEE references usually place the year near the end, not in APA-style parentheses after the author.",
                    "IEEE citation issues",
                )
            if re.search(r"^\[\d+\]\s+[A-Z][A-Za-z-]+,\s+[A-Z]\.\s*&", lines[idx]) or re.search(r"&\s+[A-Z][A-Za-z-]+,\s+[A-Z]\.", lines[idx]):
                add_change(
                    "APA-style multiple-author reference detected",
                    "high",
                    lines[idx][:80],
                    lines[idx][:80],
                    "IEEE references should use IEEE author formatting, not APA surname-initial plus ampersand formatting.",
                    "IEEE citation issues",
                )
            if re.search(r"&", lines[idx]):
                add_change(
                    "APA-style ampersand detected in reference authors",
                    "medium",
                    lines[idx][:80],
                    lines[idx][:80],
                    "IEEE references use commas and 'and' conventions rather than APA-style ampersands.",
                    "IEEE citation issues",
                )
            stripped_ref = re.sub(r"^\[\d+\]\s*", "", lines[idx]).strip()
            has_title_quotes = bool(re.search(r"[\"â€œâ€].+[\"â€œâ€]", stripped_ref))
            looks_like_article = bool(re.search(r"\b(journal|proceedings|conference|transactions|vol\.|no\.|pp\.)\b", stripped_ref, re.IGNORECASE))
            year_present = bool(re.search(r"\b(19|20)\d{2}\b", stripped_ref))
            has_publication_detail = bool(
                re.search(r"\b(journal|conference|proceedings|transactions|vol\.|no\.|pp\.|doi|press|publisher|university)\b", stripped_ref, re.IGNORECASE)
            )
            if not has_title_quotes and looks_like_article:
                add_change(
                    "Article title may need quotation marks",
                    "medium",
                    lines[idx][:80],
                    lines[idx][:80],
                    "IEEE article titles are commonly placed in quotation marks.",
                    "IEEE citation issues",
                )
            if len(stripped_ref.split()) < 7 or not year_present or not has_publication_detail:
                add_change(
                    "Reference entry appears incomplete for IEEE",
                    "high" if len(stripped_ref.split()) < 7 else "medium",
                    lines[idx][:80],
                    lines[idx][:80],
                    "IEEE references need enough metadata: authors, title, source/publication, year, and location details such as volume/pages/DOI/URL when applicable.",
                    "IEEE citation issues",
                )
            if re.match(r"^[A-Z][A-Za-z\s]+ Organization\.", stripped_ref) or re.match(r"^World Health Organization\.", stripped_ref):
                if not year_present or not re.search(r"https?://|Accessed:|accessed", stripped_ref):
                    add_change(
                        "Organization web reference is incomplete",
                        "medium",
                        lines[idx][:80],
                        lines[idx][:80],
                        "For an online organization source, include organization, page title, site name, date if available, URL, and accessed date.",
                        "IEEE citation issues",
                    )
            if re.search(r"https?://\S+\s*$", lines[idx]) and not re.search(r"Accessed:|accessed", lines[idx]):
                add_change(
                    "Raw URL reference needs full IEEE web reference details",
                    "medium",
                    lines[idx][:80],
                    lines[idx][:80],
                    "Include organization/author, page title, site name, date if available, URL, and accessed date.",
                    "IEEE citation issues",
                )
            ref_number += 1

        if body_citations:
            first_seen: List[int] = []
            for number in body_citations:
                if number not in first_seen:
                    first_seen.append(number)
            expected = list(range(1, len(first_seen) + 1))
            if first_seen != expected:
                add_change(
                    "Citation order inconsistency",
                    "high",
                    ", ".join(f"[{n}]" for n in first_seen),
                    "Manual verification required",
                    "IEEE citation numbers should follow first appearance order without skipping earlier numbers.",
                    "citation numbering issues",
                )
            skipped = [n for n in range(1, max(body_citations) + 1) if n not in body_citation_set]
            if skipped:
                add_change(
                    "Citation number skipped in body text",
                    "high",
                    ", ".join(f"[{n}]" for n in body_citations),
                    "Manual verification required",
                    f"The body cites up to [{max(body_citations)}] but never cites {', '.join(f'[{n}]' for n in skipped)}.",
                    "citation numbering issues",
                )

        reference_set = set(reference_numbers_seen or range(1, ref_number))
        uncited_refs = sorted(reference_set - body_citation_set)
        if uncited_refs and body_citations:
            add_change(
                "References not tied to in-text citations",
                "high",
                ", ".join(f"[{n}]" for n in uncited_refs),
                "Cite each listed reference or remove unused entries",
                "IEEE reference numbers should correspond to citations in the body.",
                "citation numbering issues",
            )
        fixed = "\n".join(lines)
    else:
        add_change(
            "Missing IEEE References section",
            "medium",
            "",
            "References",
            "A complete IEEE paper should end with a 'References' section.",
            "IEEE citation issues",
        )

    ieee_citations = [int(x) for x in re.findall(r"\[(\d+)\]", fixed)]
    if ieee_citations:
        first_seen: List[int] = []
        for number in ieee_citations:
            if number not in first_seen:
                first_seen.append(number)
        expected = list(range(1, len(first_seen) + 1))
        if first_seen != expected:
            add_change(
                "Citation order may be incorrect",
                "high",
                ", ".join(f"[{n}]" for n in first_seen),
                ", ".join(f"[{n}]" for n in expected),
                "IEEE citations should be numbered by first appearance.",
                "citation numbering issues",
            )

    unresolved_changes, transformation_count = split_applied_transformations(changes, fixed)
    cleaned_changes = clean_ieee_findings(unresolved_changes)
    high_count = sum(1 for item in cleaned_changes if item["severity"] == "high")
    medium_count = sum(1 for item in cleaned_changes if item["severity"] == "medium")
    summary = (
        f"IEEE formatter completed with {len(cleaned_changes)} grouped issue(s): "
        f"{high_count} high priority and {medium_count} medium priority. "
        f"{transformation_count} deterministic fix(es) were applied before reporting remaining issues. "
        "Raw rule hits were grouped and deduplicated into root problems."
    )
    return {
        "summary": summary,
        "fixed_text": fixed,
        "changes": cleaned_changes,
        "highlights": highlights,
        "transformations_applied": transformation_count > 0,
        "transformation_count": transformation_count,
        "source": "deterministic_ieee_formatter",
    }


def build_docx_from_text(text: str) -> BytesIO:
    doc = Document()
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.lower() == "references":
            doc.add_heading("References", level=1)
        else:
            doc.add_paragraph(stripped)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def extract_docx_text(file_bytes: bytes) -> str:
    from io import BytesIO

    try:
        doc = Document(BytesIO(file_bytes))
        parts: List[str] = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        full_text = "\n".join(parts).strip()
        if not full_text:
            raise ValueError("No readable text found in .docx file.")
        return full_text
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Failed to parse .docx: {exc}") from exc


def fallback_analysis(
    text: str,
    reason: str,
    review_mode: str = "academic",
    format_mode: str = "ieee",
) -> Dict[str, Any]:
    review_mode = normalize_review_mode(review_mode)
    format_mode = normalize_format_mode(format_mode)
    if review_mode == "academic":
        data = {
            "summary": f"General academic review completed with deterministic checks because the selected AI model returned invalid output ({reason}).",
            "structure_format_issues": [],
            "academic_quality_issues": [],
            "citation_consistency_issues": [],
            "prioritized_suggestions": [],
            "optional_rewrite_suggestions": [],
            "highlights": [],
            "reviewed_text": text,
            "source": "fallback",
            "fallback_used": True,
        }
        return merge_academic_safety_findings(data, text)
    data = {
        "summary": f"{format_mode.upper()} compliance review completed with deterministic checks because the model was unavailable ({reason}).",
        "structure_format_issues": [],
        "academic_quality_issues": [],
        "citation_consistency_issues": [],
        "prioritized_suggestions": [],
        "optional_rewrite_suggestions": [],
        "highlights": [],
        "reviewed_text": text,
        "source": "fallback",
        "fallback_used": True,
    }
    return merge_deterministic_findings(data, text, format_mode=format_mode)


def deterministic_findings(text: str, format_mode: str = "ieee") -> Dict[str, List[Dict[str, str]]]:
    format_mode = normalize_format_mode(format_mode)
    structure_issues: List[Dict[str, str]] = []
    citation_issues: List[Dict[str, str]] = []
    suggestions: List[Dict[str, str]] = []
    highlights: List[Dict[str, str]] = []

    def add_issue(
        target: List[Dict[str, str]],
        issue: str,
        evidence: str,
        recommendation: str,
        excerpt: str = "",
        severity: str = "medium",
    ) -> None:
        target.append(
            {
                "issue": issue,
                "severity": severity,
                "evidence": evidence,
                "recommendation": recommendation,
            }
        )
        if excerpt:
            highlights.append(
                {
                    "excerpt": excerpt,
                    "message": recommendation,
                    "severity": severity,
                    "category": f"{format_mode.upper()} compliance",
                }
            )

    lines = text.splitlines()
    has_references_heading = any(REFERENCE_HEADING_RE.match(line) for line in text.splitlines())
    wrong_reference_heading = re.search(r"(?im)^\s*(bibliography|works cited|reference list)\s*$", text)
    apa_citations = re.findall(r"\([A-Z][A-Za-z' -]+(?:\s+et al\.)?,\s*\d{4}[a-z]?\)", text)
    apa_citations.extend(re.findall(r"\b[A-Z][A-Za-z'-]+\s+\(\d{4}[a-z]?\)", text))
    body_lines_for_citations: List[str] = []
    in_refs_for_citations = False
    for line in lines:
        if REFERENCE_HEADING_RE.match(line):
            in_refs_for_citations = True
            continue
        if not in_refs_for_citations:
            body_lines_for_citations.append(line)
    body_text_for_citations = "\n".join(body_lines_for_citations)
    ieee_citations = re.findall(r"\[(\d+)\]", body_text_for_citations)
    figure_lines = [line.strip() for line in lines if re.match(r"(?i)^\s*(fig(?:ure)?\.?)\s*\d+", line)]
    table_lines = [line.strip() for line in lines if re.match(r"(?i)^\s*table\s+[A-Z0-9IVX]+", line)]

    if wrong_reference_heading:
        add_issue(
            citation_issues,
            "Reference section heading is not compliant",
            f"Detected heading '{wrong_reference_heading.group(1)}'.",
            "Rename the section heading to 'References'.",
            wrong_reference_heading.group(0).strip(),
        )
    elif not has_references_heading:
        add_issue(
            citation_issues,
            "References section is not visible",
            "No standalone 'References' heading was detected.",
            "Add a References section and verify that every in-text citation has a matching entry.",
        )

    if format_mode == "apa7":
        if ieee_citations:
            marker = f"[{ieee_citations[0]}]"
            add_issue(
                citation_issues,
                "Numeric bracket citation violates APA 7",
                f"Detected IEEE-style marker '{marker}'.",
                "Use an APA 7 author-date citation such as (Smith, 2024).",
                marker,
                "high",
            )
        ampersand_citation = re.search(r"\([A-Z][A-Za-z' -]+\s+and\s+[A-Z][A-Za-z' -]+,\s*\d{4}\)", text)
        if ampersand_citation:
            add_issue(
                citation_issues,
                "Parenthetical APA 7 citation should use an ampersand",
                f"Detected '{ampersand_citation.group(0)}'.",
                "Use '&' between author names inside a parenthetical citation.",
                ampersand_citation.group(0),
            )
        reference_lines: List[str] = []
        in_references = False
        for line in lines:
            if REFERENCE_HEADING_RE.match(line):
                in_references = True
                continue
            if in_references and line.strip():
                reference_lines.append(line.strip())
        numbered_reference = next((line for line in reference_lines if re.match(r"^\[\d+\]", line)), "")
        if numbered_reference:
            add_issue(
                citation_issues,
                "Numbered reference entry violates APA 7",
                f"Detected '{numbered_reference[:100]}'.",
                "Remove bracketed reference numbers. APA 7 reference entries are alphabetized by author surname.",
                numbered_reference[:100],
                "high",
            )
        initials_first_reference = next(
            (line for line in reference_lines if re.match(r"^(?:\[\d+\]\s*)?[A-Z]\.\s+[A-Z][A-Za-z'-]+,", line)),
            "",
        )
        if initials_first_reference:
            add_issue(
                citation_issues,
                "Reference author name order does not follow APA 7",
                f"Detected '{initials_first_reference[:100]}'.",
                "Use surname first followed by initials, for example 'Smith, J.'.",
                initials_first_reference[:100],
            )
        surnames = [
            match.group(1).lower()
            for line in reference_lines
            if (match := re.match(r"^(?:\[\d+\]\s*)?([A-Z][A-Za-z'-]+),", line))
        ]
        if len(surnames) > 1 and surnames != sorted(surnames):
            add_issue(
                citation_issues,
                "APA 7 reference entries are not alphabetized",
                f"Detected leading author surnames in this order: {surnames}.",
                "Alphabetize reference entries by the surname of the first author.",
            )
        cited_surnames = {
            match.lower()
            for match in re.findall(r"\(([A-Z][A-Za-z'-]+)(?:\s+et al\.)?,\s*\d{4}[a-z]?\)", text)
        }
        missing_surnames = sorted(cited_surnames - set(surnames))
        if missing_surnames and has_references_heading:
            add_issue(
                citation_issues,
                "APA 7 in-text citation has no visible matching reference entry",
                f"No surname-first reference entry was found for: {missing_surnames}.",
                "Add the missing reference entries or correct the in-text citations.",
                severity="high",
            )
        for line in figure_lines:
            if re.match(r"(?i)^fig\.", line) or re.match(r"(?i)^figure\s+\d+\s*[:.-]", line):
                add_issue(
                    structure_issues,
                    "Figure label does not follow APA 7 naming",
                    f"Detected '{line}'.",
                    "Use 'Figure 1' on its own line, followed by the italicized figure title on the next line.",
                    line,
                )
        for line in table_lines:
            if re.match(r"(?i)^table\s+\d+\s*[:.-]", line) or re.match(r"(?i)^table\s+[IVX]+\b", line):
                add_issue(
                    structure_issues,
                    "Table label does not follow APA 7 naming",
                    f"Detected '{line}'.",
                    "Use an Arabic-number label such as 'Table 1' on its own line, followed by the italicized title.",
                    line,
                )
    else:
        if apa_citations:
            add_issue(
                citation_issues,
                "Author-date citation violates IEEE style",
                f"Detected APA-style citation '{apa_citations[0]}'.",
                "Use a numbered IEEE citation such as [1].",
                apa_citations[0],
                "high",
            )
        for line in figure_lines:
            if not re.match(r"^Fig\.\s*\d+\.\s+\S", line):
                add_issue(
                    structure_issues,
                    "Figure caption does not follow IEEE naming",
                    f"Detected '{line}'.",
                    "Use an IEEE caption such as 'Fig. 1. Caption text.' below the figure.",
                    line,
                )
        for line in table_lines:
            if not re.match(r"^TABLE\s+[IVX]+\b", line):
                add_issue(
                    structure_issues,
                    "Table caption does not follow IEEE naming",
                    f"Detected '{line}'.",
                    "Use an uppercase Roman-numeral table label such as 'TABLE I' above the table.",
                    line,
                )

        reference_numbers: List[int] = []
        in_references = False
        unnumbered_reference = ""
        for line in lines:
            if REFERENCE_HEADING_RE.match(line):
                in_references = True
                continue
            if not in_references or not line.strip():
                continue
            match = re.match(r"^\s*\[(\d+)\]", line)
            if match:
                reference_numbers.append(int(match.group(1)))
            elif not unnumbered_reference:
                unnumbered_reference = line.strip()
        if unnumbered_reference:
            add_issue(
                citation_issues,
                "IEEE reference entry is missing a bracketed number",
                f"Detected unnumbered reference entry '{unnumbered_reference[:100]}'.",
                "Number each IEEE reference entry in citation order, for example '[1] ...'.",
                unnumbered_reference[:100],
                "high",
            )
        if reference_numbers and reference_numbers != list(range(1, len(reference_numbers) + 1)):
            add_issue(
                citation_issues,
                "IEEE reference entries are not sequentially numbered",
                f"Detected reference numbers {reference_numbers}.",
                "Number references sequentially in order of first citation: [1], [2], [3], and so on.",
                severity="high",
            )
        body_citation_numbers = [int(value) for value in ieee_citations]
        if body_citation_numbers:
            first_seen: List[int] = []
            for number in body_citation_numbers:
                if number not in first_seen:
                    first_seen.append(number)
            expected = list(range(1, len(first_seen) + 1))
            if first_seen != expected:
                add_issue(
                    citation_issues,
                    "IEEE in-text citations are not numbered by first appearance",
                    f"Detected first appearances as {first_seen}. Repeated citations are allowed, but new citation numbers should appear as {expected}.",
                    "Use IEEE citation numbers in order of first appearance. Repeats like [1], [2], [1], [3] are valid; new out-of-order citations like [2], [1], [3] are not.",
                    severity="high",
                )
        missing_entries = sorted(set(int(value) for value in ieee_citations) - set(reference_numbers))
        if missing_entries and has_references_heading:
            add_issue(
                citation_issues,
                "IEEE in-text citation has no matching reference entry",
                f"No matching reference entry was found for {missing_entries}.",
                "Add the missing numbered entries or correct the in-text citation numbers.",
                severity="high",
            )

    for item in (citation_issues + structure_issues)[:5]:
        suggestions.append(
            {
                "priority": item["severity"],
                "suggestion": item["recommendation"],
                "rationale": item["evidence"],
                "expected_impact": f"Resolves a detected {format_mode.upper()} compliance violation.",
            }
        )

    return {
        "structure_format_issues": structure_issues,
        "academic_quality_issues": [],
        "citation_consistency_issues": citation_issues,
        "prioritized_suggestions": suggestions,
        "optional_rewrite_suggestions": [],
        "highlights": highlights,
    }


def merge_deterministic_findings(data: Dict[str, Any], text: str, format_mode: str = "ieee") -> Dict[str, Any]:
    findings = deterministic_findings(text, format_mode=format_mode)
    for field, items in findings.items():
        existing = data.get(field, [])
        if not isinstance(existing, list):
            existing = []
        existing_titles = {
            str(item.get("issue") or item.get("suggestion") or item.get("reason") or "").lower()
            for item in existing
            if isinstance(item, dict)
        }
        merged = []
        for item in items:
            key = str(item.get("issue") or item.get("suggestion") or item.get("reason") or "").lower()
            if key and key not in existing_titles:
                merged.append(item)
        data[field] = merged + existing
    return data


def merge_academic_safety_findings(data: Dict[str, Any], text: str) -> Dict[str, Any]:
    academic_issues = data.get("academic_quality_issues", [])
    citation_issues = data.get("citation_consistency_issues", [])
    suggestions = data.get("prioritized_suggestions", [])
    highlights = data.get("highlights", [])
    if not isinstance(academic_issues, list):
        academic_issues = []
    if not isinstance(citation_issues, list):
        citation_issues = []
    if not isinstance(suggestions, list):
        suggestions = []
    if not isinstance(highlights, list):
        highlights = []

    existing_titles = {
        str(item.get("issue") or item.get("suggestion") or "").strip().lower()
        for item in [*academic_issues, *citation_issues, *suggestions]
        if isinstance(item, dict)
    }

    def add_academic_issue(issue: str, severity: str, evidence: str, recommendation: str, excerpt: str = "") -> None:
        key = issue.lower()
        if key in existing_titles:
            return
        academic_issues.append(
            {
                "issue": issue,
                "severity": severity,
                "evidence": evidence,
                "recommendation": recommendation,
            }
        )
        suggestions.append(
            {
                "priority": severity,
                "suggestion": recommendation,
                "rationale": evidence,
                "expected_impact": "Improves the paper's academic focus, depth, or scholarly credibility.",
            }
        )
        if excerpt:
            highlights.append(
                {
                    "excerpt": excerpt,
                    "message": f"{issue}: {recommendation}",
                    "severity": severity,
                    "category": "academic quality",
                }
            )
        existing_titles.add(key)

    def add_citation_issue(issue: str, severity: str, evidence: str, recommendation: str, excerpt: str = "") -> None:
        key = issue.lower()
        if key in existing_titles:
            return
        citation_issues.insert(
            0,
            {
                "issue": issue,
                "severity": severity,
                "evidence": evidence,
                "recommendation": recommendation,
            },
        )
        suggestions.insert(
            0,
            {
                "priority": severity,
                "suggestion": recommendation,
                "rationale": evidence,
                "expected_impact": "Prevents the review from accepting unsupported or unverifiable source markers as evidence.",
            },
        )
        if excerpt:
            highlights.append(
                {
                    "excerpt": excerpt,
                    "message": f"{issue}: {recommendation}",
                    "severity": severity,
                    "category": "citation verification",
                }
            )
        existing_titles.add(key)

    lines = text.splitlines()
    body_lines_for_citations: List[str] = []
    reference_lines: List[str] = []
    in_references_for_citations = False
    for line in lines:
        if REFERENCE_HEADING_RE.match(line):
            in_references_for_citations = True
            continue
        if in_references_for_citations:
            reference_lines.append(line)
        else:
            body_lines_for_citations.append(line)

    body_for_citations = "\n".join(body_lines_for_citations)
    body_bracket_numbers = [int(value) for value in re.findall(r"\[(\d+)\]", body_for_citations)]
    bracket_numbers = body_bracket_numbers or [int(value) for value in re.findall(r"\[(\d+)\]", text)]
    has_references_heading = any(REFERENCE_HEADING_RE.match(line) for line in text.splitlines())
    if bracket_numbers and not has_references_heading:
        add_citation_issue(
            "Bracket citations have no visible reference list",
            "high",
            f"Detected bracket citations {sorted(set(bracket_numbers))}, but no standalone References section was found.",
            "Add complete reference entries for every bracket citation and verify that the cited sources actually exist.",
            f"[{bracket_numbers[0]}]",
        )

    if has_references_heading:
        reference_numbers = set()
        reference_number_order: List[int] = []
        incomplete_ref = ""
        for line in reference_lines:
            if not line.strip():
                continue
            ref_match = re.match(r"^\s*\[(\d+)\]\s+(.+)$", line.strip())
            if ref_match:
                reference_number = int(ref_match.group(1))
                reference_numbers.add(reference_number)
                reference_number_order.append(reference_number)
                body = ref_match.group(2)
                if not incomplete_ref and (len(body.split()) < 6 or not re.search(r"\b(19|20)\d{2}\b", body)):
                    incomplete_ref = line.strip()
        if reference_number_order and reference_number_order != list(range(1, len(reference_number_order) + 1)):
            add_citation_issue(
                "Reference list numbering is not sequential",
                "high",
                f"Detected reference entries numbered {reference_number_order}.",
                "Renumber the References section sequentially as [1], [2], [3], and so on, matching first citation order.",
                f"[{reference_number_order[-1]}]",
            )
        missing = sorted(set(bracket_numbers) - reference_numbers)
        if missing:
            add_citation_issue(
                "Bracket citations lack matching reference entries",
                "high",
                f"In-text citation numbers {missing} do not have matching numbered entries in the References section.",
                "Add the missing source metadata or remove unsupported citation markers.",
                f"[{missing[0]}]",
            )
        if incomplete_ref:
            add_citation_issue(
                "Reference entry appears incomplete or placeholder-like",
                "high",
                f"Detected incomplete reference entry: {incomplete_ref[:160]}",
                "Verify the source exists and provide complete author, title, publication, date, and retrieval details as required.",
                incomplete_ref[:160],
            )

    if body_bracket_numbers:
        first_seen: List[int] = []
        for number in body_bracket_numbers:
            if number not in first_seen:
                first_seen.append(number)
        expected = list(range(1, len(first_seen) + 1))
        if first_seen != expected:
            add_citation_issue(
                "In-text citation numbers are not in first-appearance order",
                "high",
                f"Detected first appearances as {first_seen}. Repeated citations are allowed, but new citation numbers should appear as {expected}.",
                "Use IEEE citation numbers in order of first appearance. Repeats like [1], [2], [1], [3] are valid; new out-of-order citations like [2], [1], [3] are not.",
                f"[{first_seen[0]}]",
            )

    body_text = []
    for line in text.splitlines():
        if REFERENCE_HEADING_RE.match(line):
            break
        body_text.append(line)
    body_for_claim_checks = "\n".join(body_text)

    claim_pattern = re.compile(
        r"[^.!?]*(?:improves?|enhances?|increases?|reduces?|causes?|leads to|beneficial tool|cognitive performance|academic performance)[^.!?]*[.!?]",
        re.IGNORECASE,
    )
    citation_like = re.compile(r"\[\d+\]|\([A-Z][A-Za-z' -]+,\s*\d{4}[a-z]?\)")
    for match in claim_pattern.finditer(body_for_claim_checks):
        sentence = match.group(0).strip()
        if sentence and not citation_like.search(sentence):
            add_citation_issue(
                "Evidence-based claim may need citation",
                "high",
                f"Detected an externally verifiable claim without a visible citation: {sentence[:180]}",
                "Add a credible source or revise the sentence as an unsupported claim.",
                sentence[:180],
            )
            break

    abstract_match = re.search(
        r"(?is)(?:^|\n)\s*#{0,6}\s*abstract\s*\n+(.+?)(?=\n\s*#{0,6}\s*(?:[IVX]+\.\s*)?introduction\b|\Z)",
        body_for_claim_checks,
    )
    if abstract_match:
        abstract_text = " ".join(abstract_match.group(1).split())
        if re.search(r"\bbrief overview\b|\bprovides an overview\b|\btalks about\b|\bcontinued significance\b|\bvery old\b", abstract_text, re.IGNORECASE):
            add_academic_issue(
                "Abstract is generic",
                "medium",
                f"Detected a broad descriptive abstract: {abstract_text[:180]}",
                "State the paper's specific focus, scope, and main takeaway instead of only saying it provides an overview.",
                abstract_text[:180],
            )

    objective_match = re.search(r"(?i)\b(this paper|this study)\s+(?:provides|discusses|examines|explores|talks about)\b[^.!?]*[.!?]", body_for_claim_checks)
    if objective_match:
        add_academic_issue(
            "Weak thesis or research contribution",
            "high",
            f"Detected a descriptive objective without a clear argument or contribution: {objective_match.group(0).strip()}",
            "Add a focused thesis, research question, or contribution that explains what the paper adds beyond summary.",
            objective_match.group(0).strip(),
        )

    body_sections = re.findall(r"(?im)^\s*[IVX]+\.\s+.+$", body_for_claim_checks)
    if len(body_sections) >= 3 and not re.search(r"(?i)\btherefore|however|in contrast|this suggests|this demonstrates|because|as a result\b", body_for_claim_checks):
        add_academic_issue(
            "Writing is mostly descriptive",
            "medium",
            "The draft is organized as a sequence of historical facts but shows little analysis, synthesis, or argument.",
            "Add analytical transitions that explain why each historical stage matters for the paper's central claim.",
        )

    reference_years = [int(year) for year in re.findall(r"(?m)^\s*\[\d+\].*?\b((?:18|19|20)\d{2})\b", text)]
    old_reference_years = [year for year in reference_years if year < 1990]
    if len(old_reference_years) >= max(2, len(reference_years) // 2):
        add_academic_issue(
            "Source base may be dated",
            "medium",
            f"Several references are older sources: {', '.join(str(year) for year in old_reference_years[:6])}.",
            "Keep historically important sources, but add recent scholarship if the paper makes claims about current significance or modern chess.",
        )

    data["academic_quality_issues"] = academic_issues
    data["citation_consistency_issues"] = citation_issues
    data["prioritized_suggestions"] = suggestions
    data["highlights"] = highlights
    return data


def coerce_analysis_shape(
    data: Dict[str, Any],
    text: str,
    review_mode: str = "academic",
    format_mode: str = "ieee",
) -> Dict[str, Any]:
    review_mode = normalize_review_mode(review_mode)
    required_list_fields = [
        "structure_format_issues",
        "academic_quality_issues",
        "citation_consistency_issues",
        "prioritized_suggestions",
        "optional_rewrite_suggestions",
        "highlights",
    ]
    if "summary" not in data or not isinstance(data["summary"], str):
        data["summary"] = "Automated analysis completed."
    for field in required_list_fields:
        if field not in data or not isinstance(data[field], list):
            data[field] = []
    issue_fields = [
        "structure_format_issues",
        "academic_quality_issues",
        "citation_consistency_issues",
    ]
    for field in issue_fields:
        data[field] = [
            item
            for item in data[field]
            if isinstance(item, dict)
            and str(item.get("issue", "")).strip()
            and str(item.get("evidence") or item.get("recommendation") or "").strip()
        ]
    data["prioritized_suggestions"] = [
        item
        for item in data["prioritized_suggestions"]
        if isinstance(item, dict)
        and str(item.get("suggestion", "")).strip()
        and str(item.get("rationale") or item.get("expected_impact") or "").strip()
    ]
    data["optional_rewrite_suggestions"] = [
        item
        for item in data["optional_rewrite_suggestions"]
        if isinstance(item, dict)
        and str(item.get("original_excerpt", "")).strip()
        and str(item.get("rewritten_excerpt", "")).strip()
    ]
    data["highlights"] = [
        item
        for item in data["highlights"]
        if isinstance(item, dict)
        and str(item.get("excerpt", "")).strip()
        and str(item.get("message", "")).strip()
    ]
    existing_highlight_keys = {str(item.get("excerpt", "")).strip().lower() for item in data["highlights"]}
    for field in issue_fields:
        for item in data[field]:
            evidence = str(item.get("evidence", "")).strip()
            issue = str(item.get("issue", "Review this excerpt")).strip()
            recommendation = str(item.get("recommendation", "")).strip()
            severity = str(item.get("severity", "medium")).strip() or "medium"
            if not evidence or evidence.lower().startswith("no precise evidence"):
                continue
            excerpt = evidence
            if len(excerpt) > 240:
                sentence_match = re.search(r"[^.!?]{20,220}[.!?]", excerpt)
                excerpt = sentence_match.group(0).strip() if sentence_match else excerpt[:240].strip()
            key = excerpt.lower()
            if key in existing_highlight_keys:
                continue
            data["highlights"].append(
                {
                    "excerpt": excerpt,
                    "message": f"{issue}: {recommendation}" if recommendation else issue,
                    "severity": severity,
                    "category": issue,
                }
            )
            existing_highlight_keys.add(key)
    data["source"] = "openai"
    data["fallback_used"] = False
    data["reviewed_text"] = text

    if review_mode == "format":
        findings = deterministic_findings(text, format_mode=format_mode)
        violation_count = (
            len(findings["structure_format_issues"])
            + len(findings["citation_consistency_issues"])
        )
        style_label = {"apa7": "APA 7", "ieee": "IEEE"}[normalize_format_mode(format_mode)]
        data["summary"] = (
            f"{style_label} format check found {violation_count} visible violation(s)."
            if violation_count
            else f"No visible {style_label} formatting or citation violations were detected by the rule-based checker."
        )
        data["structure_format_issues"] = findings["structure_format_issues"]
        data["academic_quality_issues"] = []
        data["citation_consistency_issues"] = findings["citation_consistency_issues"]
        data["prioritized_suggestions"] = findings["prioritized_suggestions"]
        data["optional_rewrite_suggestions"] = []
        data["highlights"] = findings["highlights"]
        return data
    return merge_academic_safety_findings(data, text)


def build_analysis_prompt(
    text: str,
    metadata: Optional[Dict[str, Any]] = None,
    review_mode: str = "academic",
    format_mode: str = "ieee",
) -> str:
    review_mode = normalize_review_mode(review_mode)
    if review_mode == "academic":
        return f"""
You are a rigorous academic writing reviewer. Review the submitted paper for writing quality and academic effectiveness.
Return ONLY valid JSON, no markdown.
JSON keys required:
- summary (string)
- structure_format_issues (array of objects: issue, severity, evidence, recommendation)
- academic_quality_issues (array of objects: issue, severity, evidence, recommendation)
- citation_consistency_issues (array of objects: issue, severity, evidence, recommendation)
- prioritized_suggestions (array of objects: priority, suggestion, rationale, expected_impact)
- optional_rewrite_suggestions (array of objects: original_excerpt, rewritten_excerpt, reason)
- highlights (array of objects: excerpt, message, severity, category)

Review priorities:
- Do not look for errors just to fill the output. If there is no meaningful issue in a category, return an empty array.
- Prioritize major academic risks over minor rewriting: unsupported claims, fake or unverifiable citations, missing reference entries, weak thesis/gap/objective, evidence mismatch, and unclear logic.
- Treat bracket citations such as [1], [2], [3], [4] as unverified source markers unless the draft includes matching complete reference entries. Question whether the cited sources actually exist when metadata is missing, incomplete, or placeholder-like.
- Flag claims that appear to need citations, including factual assertions, causal claims, statistics, attributed ideas, and broad research claims such as claims that chess improves cognitive performance.
- For survey, history, overview, or background papers, check whether the paper has a clear thesis, research question, contribution, scope, and synthesis rather than only encyclopedic summary.
- Flag a generic abstract when it only says the paper provides a brief overview without a specific argument, finding, scope, or takeaway.
- Flag dated source bases when the paper relies heavily on old sources for current claims, while recognizing that old sources can still be valid primary or historical sources.
- For introductions, judge only whether they provide enough background, establish a gap/problem, and state a clear objective. Do not demand detailed analysis inside the introduction.
- Explain how the paper sounds only when there is a real clarity, tone, flow, grammar, concision, or precision problem.
- Do not criticize acceptable academic wording merely because it could be rephrased. Avoid low-value comments about word choice when the phrase is clear and appropriate.
- Never cite or quote a sentence that is not actually present in the submitted paper.
- Suggest rewrites only for high-impact examples tied to a real issue.
- Every issue must include exact concrete evidence copied from the paper.
- Add a highlight for each meaningful issue using the exact shortest excerpt from the paper.
- If multiple issues refer to the same phrase or sentence, reuse the same excerpt with a different message.
- Do not invent sources, facts, quotations, or paper requirements.
- Priority ranking should put citation/source validity and unsupported major claims before minor flow, wording, or style comments.
- Prefer fewer useful findings over generic filler. Keep the summary under 120 words.

Context metadata:
{json.dumps(metadata or {}, ensure_ascii=True)}

Text to review:
{text}
"""
    style_label = {"apa7": "APA 7", "ieee": "IEEE"}[format_mode]
    return f"""
You are a strict {style_label} document-compliance reviewer.
Analyze only violations of the selected style. Do not critique the argument, research quality, methodology, or writing quality.
Return ONLY valid JSON, no markdown.
JSON keys required:
- summary (string)
- structure_format_issues (array of objects: issue, severity, evidence, recommendation)
- academic_quality_issues (array of objects: issue, severity, evidence, recommendation)
- citation_consistency_issues (array of objects: issue, severity, evidence, recommendation)
- prioritized_suggestions (array of objects: priority, suggestion, rationale, expected_impact)
- optional_rewrite_suggestions (array of objects: original_excerpt, rewritten_excerpt, reason)
- highlights (array of objects: excerpt, message, severity, category)

Constraints:
- Report only concrete {style_label} violations visible in the submitted text.
- Check formatting conventions, in-text citations, references, citation/reference matching, table naming, and figure naming.
- Do not invent issues just to fill categories.
- Every issue must cite concrete evidence from the draft.
- Add a highlight for each meaningful issue using the exact shortest excerpt from the draft.
- If multiple issues refer to the same phrase or sentence, reuse the same excerpt with a different message.
- If a category has no meaningful issue, return an empty array for that category.
- academic_quality_issues must always be an empty array.
- optional_rewrite_suggestions must always be an empty array.
- Severity must reflect compliance risk: high, medium, or low.
- Keep summary under 100 words.

Style requirements:
- APA 7 uses author-date in-text citations, the heading "References", Arabic-number labels such as "Table 1" and "Figure 1", and APA table/figure title placement.
- IEEE uses bracketed numeric in-text citations such as [1], references numbered in citation order, the heading "References", figure captions such as "Fig. 1. Caption", and uppercase Roman-numeral table labels such as "TABLE I".
- Flag mixed citation systems and malformed or unmatched references when visible.

Output expectations:
- structure_format_issues should contain formatting, table, and figure violations.
- citation_consistency_issues should contain citation and reference violations.
- prioritized_suggestions should rank the most important compliance fixes.
- highlights should contain exact short excerpts copied from the text that deserve inline marking.
- highlight messages should be concise hover text.

Context metadata:
{json.dumps(metadata or {}, ensure_ascii=True)}

Text to review:
{text}
"""


def build_ollama_analysis_prompt(
    text: str,
    metadata: Optional[Dict[str, Any]] = None,
    review_mode: str = "academic",
    format_mode: str = "ieee",
) -> str:
    review_mode = normalize_review_mode(review_mode)
    if review_mode != "academic":
        return build_analysis_prompt(text=text, metadata=metadata, review_mode=review_mode, format_mode=format_mode)
    return f"""
Return one valid compact JSON object only. No markdown. No explanation outside JSON.

Required JSON shape:
{{
  "summary": "one sentence under 45 words",
  "structure_format_issues": [],
  "academic_quality_issues": [
    {{"issue":"", "severity":"high|medium|low", "evidence":"exact short quote from paper", "recommendation":""}}
  ],
  "citation_consistency_issues": [
    {{"issue":"", "severity":"high|medium|low", "evidence":"exact short quote from paper", "recommendation":""}}
  ],
  "prioritized_suggestions": [
    {{"priority":"high|medium|low", "suggestion":"", "rationale":"", "expected_impact":""}}
  ],
  "optional_rewrite_suggestions": [],
  "highlights": []
}}

Rules:
- Return at most 4 academic_quality_issues, at most 3 citation_consistency_issues, and at most 5 prioritized_suggestions.
- Leave optional_rewrite_suggestions and highlights as empty arrays.
- Do not invent issues. If no real issue exists, use an empty array.
- Focus on major academic problems: generic abstract, weak thesis or contribution, descriptive/encyclopedic writing, unsupported claims, missing or mismatched references, dated source base, grammar/tone only if clearly weak.
- Do not say IEEE citation numbering is inconsistent when citations are [1], [2], [3] in first-appearance order.
- Every issue must use exact evidence that appears in the paper.

Metadata:
{json.dumps(metadata or {}, ensure_ascii=True)}

Paper:
{text}
"""


def parse_model_analysis(
    raw: str,
    text: str,
    review_mode: str = "academic",
    format_mode: str = "ieee",
) -> Dict[str, Any]:
    if not raw:
        raise ValueError("Empty response from selected model.")
    data = json.loads(raw)
    if not isinstance(data, dict):
        raise ValueError("Selected model output is not a JSON object.")
    return coerce_analysis_shape(data, text, review_mode=review_mode, format_mode=format_mode)


def json_line(event: str, payload: Dict[str, Any]) -> str:
    return json.dumps({"event": event, **payload}, ensure_ascii=True) + "\n"


def remember_paper_context(
    session_id: Optional[str],
    text: str,
    analysis: Dict[str, Any],
    review_mode: str,
    format_mode: str,
) -> None:
    if not session_id:
        return
    state = chat_sessions.setdefault(session_id, {"phase": "chat", "history": []})
    issue_context = json.dumps(
        {
            "structure_or_writing_issues": analysis.get("structure_format_issues", [])[:8],
            "academic_quality_issues": analysis.get("academic_quality_issues", [])[:8],
            "citation_issues": analysis.get("citation_consistency_issues", [])[:8],
            "priority_fixes": analysis.get("prioritized_suggestions", [])[:8],
        },
        ensure_ascii=True,
    )
    state["paper_context"] = (
        f"Review mode: {review_mode}. Format style: {format_mode}.\n"
        f"Review summary: {analysis.get('summary', '')}\n"
        f"Review findings: {issue_context}\n"
        f"Paper text:\n{text[:12000]}"
    )
    state["paper_sections"] = extract_paper_sections(text)



@app.get("/health")
def health() -> Dict[str, Any]:
    ollama_reachable = True
    try:
        requests.get(f"{get_ollama_url()}/api/tags", timeout=3).raise_for_status()
    except Exception:
        ollama_reachable = False
    return {
        "status": "ok",
        "service": "academic-review-mvp",
        "analysis_model": get_analysis_model(),
        "openai_configured": bool(os.getenv("OPENAI_API_KEY", "").strip()),
        "ollama_reachable": ollama_reachable,
    }


@app.post("/chat", response_model=ChatResponse)
def chat(req: ChatRequest) -> ChatResponse:
    state = chat_sessions.setdefault(
        req.session_id,
        {"phase": "chat", "history": []},
    )
    message = req.message.strip()
    state["history"].append({"role": "user", "content": message})
    state["history"] = state["history"][-12:]
    focused_section = find_requested_section(message, state.get("paper_sections", []))
    try:
        reply = call_assistant_chat(
            state["history"],
            model_override=req.model,
            paper_context=state.get("paper_context"),
            focused_section=focused_section,
        )
    except Exception:
        state["history"].pop()
        raise
    state["history"].append({"role": "assistant", "content": reply})
    return ChatResponse(
        session_id=req.session_id,
        phase="chat",
        questions=[],
        next_prompt=reply,
        context={
            "paper_context_available": bool(state.get("paper_context")),
            "focused_section": focused_section.get("title") if focused_section else None,
            "focused_section_found": bool(focused_section and focused_section.get("content")),
            "focused_section_text": focused_section.get("content", "") if focused_section else "",
        },
    )


@app.post("/analyze_text", response_model=AnalysisResponse)
def analyze_text(req: AnalyzeTextRequest) -> AnalysisResponse:
    text = req.text.strip()
    if len(text) < 20:
        raise HTTPException(status_code=422, detail="Text too short for meaningful analysis.")
    data = analyze_text_with_resilience(
        text=text,
        metadata=req.metadata,
        model_override=req.model,
        review_mode=req.review_mode or "academic",
        format_mode=req.format_mode or "ieee",
    )
    remember_paper_context(req.session_id, text, data, req.review_mode or "academic", req.format_mode or "ieee")
    return AnalysisResponse(**data)


@app.post("/analyze_text_stream")
def analyze_text_stream(req: AnalyzeTextStreamRequest) -> StreamingResponse:
    return StreamingResponse(
        analysis_event_stream(
            text=req.text.strip(),
            metadata=req.metadata,
            model_override=req.model,
            session_id=req.session_id,
            review_mode=req.review_mode or "academic",
            format_mode=req.format_mode or "ieee",
        ),
        media_type="application/x-ndjson",
    )


@app.post("/analyze_docx", response_model=AnalysisResponse)
async def analyze_docx(
    file: UploadFile = File(...),
    session_id: Optional[str] = Form(None),
    model: Optional[str] = Form(None),
    review_mode: str = Form("academic"),
    format_mode: str = Form("ieee"),
) -> AnalysisResponse:
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported.")
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")
    text = extract_docx_text(content)
    metadata = {"filename": file.filename, "source_type": "docx_upload"}
    data = analyze_text_with_resilience(
        text=text,
        metadata=metadata,
        model_override=model,
        review_mode=review_mode,
        format_mode=format_mode,
    )
    remember_paper_context(session_id, text, data, review_mode, format_mode)
    return AnalysisResponse(**data)


@app.post("/analyze_docx_stream")
async def analyze_docx_stream(
    file: UploadFile = File(...),
    session_id: Optional[str] = Form(None),
    model: Optional[str] = Form(None),
    review_mode: str = Form("academic"),
    format_mode: str = Form("ieee"),
) -> StreamingResponse:
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported.")
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    def stream() -> Generator[str, None, None]:
        yield json_line(
            "status",
            {
                "step": "Parsing DOCX",
                "detail": f"Extracting readable text from {file.filename}.",
            },
        )
        try:
            text = extract_docx_text(content)
        except HTTPException as exc:
            yield json_line("error", {"detail": exc.detail, "error_type": exc.__class__.__name__})
            return
        metadata = {"filename": file.filename, "source_type": "docx_upload"}
        yield from analysis_event_stream(
            text=text,
            metadata=metadata,
            model_override=model,
            session_id=session_id,
            review_mode=review_mode,
            format_mode=format_mode,
        )

    return StreamingResponse(stream(), media_type="application/x-ndjson")


@app.post("/format_ieee_text", response_model=FormatResponse)
def format_ieee_text(req: FormatTextRequest) -> FormatResponse:
    return FormatResponse(**ieee_format_text(req.text.strip()))


@app.post("/format_ieee_docx", response_model=FormatResponse)
async def format_ieee_docx(file: UploadFile = File(...)) -> FormatResponse:
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported.")
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")
    text = extract_docx_text(content)
    return FormatResponse(**ieee_format_text(text))


@app.post("/export_ieee_docx")
def export_ieee_docx(req: ExportDocxRequest) -> StreamingResponse:
    filename = req.filename if req.filename.lower().endswith(".docx") else f"{req.filename}.docx"
    safe_filename = re.sub(r"[^A-Za-z0-9_.-]+", "_", filename)
    output = build_docx_from_text(req.text)
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_filename}"'},
    )


@app.post("/format_ieee_docx_export")
async def format_ieee_docx_export(file: UploadFile = File(...)) -> StreamingResponse:
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported.")
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")
    text = extract_docx_text(content)
    formatted = ieee_format_text(text)["fixed_text"]
    base = re.sub(r"\.docx$", "", file.filename, flags=re.IGNORECASE)
    safe_filename = re.sub(r"[^A-Za-z0-9_.-]+", "_", f"{base}_ieee_formatted.docx")
    output = build_docx_from_text(formatted)
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_filename}"'},
    )


# Provider-aware overrides. These are intentionally defined after the routes because
# route handlers resolve these globals at request time.
def call_model_for_analysis(
    text: str,
    metadata: Optional[Dict[str, Any]] = None,
    model_override: Optional[str] = None,
    review_mode: str = "academic",
    format_mode: str = "ieee",
) -> Dict[str, Any]:
    model = get_analysis_model(model_override)
    review_mode = normalize_review_mode(review_mode)
    format_mode = normalize_format_mode(format_mode)
    if review_mode == "format":
        analysis = coerce_analysis_shape(
            {"summary": "Rule-based format check completed."},
            text,
            review_mode=review_mode,
            format_mode=format_mode,
        )
        analysis["source"] = "rules"
        return analysis

    if not is_openai_model(model):
        prompt = build_ollama_analysis_prompt(text=text, metadata=metadata, review_mode=review_mode, format_mode=format_mode)
        resp = requests.post(
            f"{get_ollama_url()}/api/generate",
            json={
                "model": model,
                "prompt": prompt,
                "stream": False,
                "format": "json",
                "options": ollama_options("analysis"),
            },
            timeout=get_ollama_timeout(),
        )
        if resp.status_code >= 400:
            raise RuntimeError(f"Ollama request failed with HTTP {resp.status_code}: {resp.text.strip()}")
        raw = resp.json().get("response", "").strip()
        analysis = parse_model_analysis(raw=raw, text=text, review_mode=review_mode, format_mode=format_mode)
        analysis["source"] = "ollama"
        return analysis

    prompt = build_analysis_prompt(text=text, metadata=metadata, review_mode=review_mode, format_mode=format_mode)
    body = {
        "model": openai_model_id(model),
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.2,
        "response_format": {"type": "json_object"},
        "max_tokens": 2500,
    }
    resp = requests.post(
        "https://api.openai.com/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {get_openai_api_key()}",
            "Content-Type": "application/json",
        },
        json=body,
        timeout=90,
    )
    if resp.status_code >= 400:
        raise RuntimeError(f"OpenAI request failed with HTTP {resp.status_code}: {resp.text.strip()}")
    payload = resp.json()
    raw = payload["choices"][0]["message"]["content"].strip()
    return parse_model_analysis(raw=raw, text=text, review_mode=review_mode, format_mode=format_mode)


def analyze_text_with_resilience(
    text: str,
    metadata: Optional[Dict[str, Any]] = None,
    model_override: Optional[str] = None,
    review_mode: str = "academic",
    format_mode: str = "ieee",
) -> Dict[str, Any]:
    try:
        return call_model_for_analysis(
            text=text,
            metadata=metadata,
            model_override=model_override,
            review_mode=review_mode,
            format_mode=format_mode,
        )
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Hosted/local model analysis failed")
        return fallback_analysis(text=text, reason=str(exc), review_mode=review_mode, format_mode=format_mode)


def analysis_event_stream(
    text: str,
    metadata: Optional[Dict[str, Any]] = None,
    model_override: Optional[str] = None,
    session_id: Optional[str] = None,
    review_mode: str = "academic",
    format_mode: str = "ieee",
) -> Generator[str, None, None]:
    try:
        model = get_analysis_model(model_override)
        review_mode = normalize_review_mode(review_mode)
        format_mode = normalize_format_mode(format_mode)
        yield json_line("status", {"step": "Validating input", "detail": "Checking draft length and selected model."})
        if len(text.strip()) < 20:
            yield json_line("error", {"detail": "Text too short for meaningful analysis."})
            return

        if review_mode == "format":
            yield json_line(
                "status",
                {
                    "step": "Running rule-based format checker",
                    "detail": "Checking visible APA/IEEE formatting and citation rules without contacting an AI model.",
                },
            )
            analysis = coerce_analysis_shape(
                {"summary": "Rule-based format check completed."},
                text,
                review_mode=review_mode,
                format_mode=format_mode,
            )
            analysis["source"] = "rules"
            remember_paper_context(session_id, text, analysis, review_mode, format_mode)
            yield json_line("analysis", {"analysis": analysis})
            return

        provider = "OpenAI" if is_openai_model(model) else "Ollama"
        yield json_line(
            "status",
            {
                "step": "Preparing prompt",
                "detail": f"Building the {review_mode} review request for {provider}.",
            },
        )
        yield json_line("status", {"step": f"Contacting {provider}", "detail": f"Sending the draft to {model}."})

        raw_parts: List[str] = []
        chunk_count = 0

        if is_openai_model(model):
            url = "https://api.openai.com/v1/chat/completions"
            headers = {
                "Authorization": f"Bearer {get_openai_api_key()}",
                "Content-Type": "application/json",
            }
            body = {
                "model": openai_model_id(model),
                "messages": [{"role": "user", "content": build_analysis_prompt(text=text, metadata=metadata, review_mode=review_mode, format_mode=format_mode)}],
                "temperature": 0.2,
                "response_format": {"type": "json_object"},
                "max_tokens": 2500,
                "stream": True,
            }
        else:
            url = f"{get_ollama_url()}/api/generate"
            headers = {}
            body = {
                "model": model,
                "prompt": build_ollama_analysis_prompt(text=text, metadata=metadata, review_mode=review_mode, format_mode=format_mode),
                "stream": True,
                "format": "json",
                "options": ollama_options("analysis"),
            }
        request_timeout = 90 if is_openai_model(model) else get_ollama_timeout()
        with requests.post(url, headers=headers, json=body, timeout=request_timeout, stream=True) as resp:
            if resp.status_code >= 400:
                raise RuntimeError(f"{provider} request failed with HTTP {resp.status_code}: {resp.text.strip()}")
            for line in resp.iter_lines(decode_unicode=True):
                if not line:
                    continue
                if is_openai_model(model):
                    if not line.startswith("data: "):
                        continue
                    data = line.removeprefix("data: ").strip()
                    if data == "[DONE]":
                        break
                    payload = json.loads(data)
                    piece = payload.get("choices", [{}])[0].get("delta", {}).get("content", "")
                else:
                    payload = json.loads(line)
                    piece = payload.get("response", "")
                if piece:
                    raw_parts.append(piece)
                    chunk_count += 1
                    if chunk_count == 1 or chunk_count % 100 == 0:
                        yield json_line(
                            "status",
                            {
                                "step": f"Receiving {provider} output",
                                "detail": f"{provider} has streamed {chunk_count} response chunks.",
                            },
                        )
                    if not is_openai_model(model) and chunk_count > 3000:
                        raise RuntimeError(
                            "Ollama exceeded the local output limit before completing a valid analysis."
                        )
                if not is_openai_model(model) and payload.get("done"):
                    break

        yield json_line("status", {"step": "Parsing model output", "detail": "Normalizing the structured feedback response."})
        raw = "".join(raw_parts).strip()
        if is_openai_model(model):
            analysis = parse_model_analysis(
                raw=raw,
                text=text,
                review_mode=review_mode,
                format_mode=format_mode,
            )
        else:
            try:
                analysis = parse_model_analysis(
                    raw=raw,
                    text=text,
                    review_mode=review_mode,
                    format_mode=format_mode,
                )
            except json.JSONDecodeError as exc:
                yield json_line(
                    "status",
                    {
                        "step": "Using deterministic academic fallback",
                        "detail": "Gemma returned malformed JSON, so the backend is using rule-based academic checks.",
                    },
                )
                analysis = fallback_analysis(
                    text=text,
                    reason=f"JSONDecodeError: {exc}",
                    review_mode=review_mode,
                    format_mode=format_mode,
                )
        analysis["source"] = "openai" if is_openai_model(model) else "ollama"
        remember_paper_context(session_id, text, analysis, review_mode, format_mode)
        yield json_line("analysis", {"analysis": analysis})
    except HTTPException as exc:
        yield json_line("error", {"detail": exc.detail, "error_type": exc.__class__.__name__})
    except Exception as exc:
        logger.exception("Streaming analysis failed")
        yield json_line(
            "status",
            {
                "step": "Using fallback",
                "detail": "The selected model failed, so the backend is generating a safe fallback review.",
            },
        )
        analysis = fallback_analysis(
            text=text,
            reason=f"{exc.__class__.__name__}: {exc}",
            review_mode=review_mode,
            format_mode=format_mode,
        )
        remember_paper_context(session_id, text, analysis, review_mode, format_mode)
        yield json_line("analysis", {"analysis": analysis})
